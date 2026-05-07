"""
CV Parsing Service
Handles CV text extraction and processing
"""

import logging
from typing import Dict, Optional, Tuple
import PyPDF2
from docx import Document

logger = logging.getLogger(__name__)


class CVParser:
    """
    CV Parser Service
    Extracts text from PDF, DOC, and DOCX files
    """
    
    @staticmethod
    def extract_text_from_pdf(file_path: str) -> Tuple[str, bool]:
        """
        Extract text from PDF file
        Returns: (text, success)
        """
        try:
            text = []
            with open(file_path, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)
                
                # Check if PDF is image-based (no extractable text)
                if len(pdf_reader.pages) == 0:
                    return "", False
                
                for page_num, page in enumerate(pdf_reader.pages):
                    try:
                        page_text = page.extract_text()
                        if page_text:
                            text.append(page_text)
                    except Exception as e:
                        logger.warning(f"Error extracting text from PDF page {page_num}: {e}")
                        continue
            
            extracted_text = "\n".join(text)
            
            # Check if text was successfully extracted
            if not extracted_text or len(extracted_text.strip()) < 10:
                logger.warning("PDF appears to be image-based or empty")
                return extracted_text, len(extracted_text.strip()) > 10
            
            return extracted_text, True
        
        except Exception as e:
            logger.error(f"Error extracting text from PDF: {e}")
            return "", False
    
    @staticmethod
    def extract_text_from_docx(file_path: str) -> Tuple[str, bool]:
        """
        Extract text from DOCX file
        Returns: (text, success)
        """
        try:
            doc = Document(file_path)
            paragraphs = [para.text for para in doc.paragraphs if para.text.strip()]
            
            # Extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        if cell.text.strip():
                            paragraphs.append(cell.text)
            
            extracted_text = "\n".join(paragraphs)
            
            if not extracted_text or len(extracted_text.strip()) < 10:
                logger.warning("DOCX file is empty or has insufficient content")
                return extracted_text, len(extracted_text.strip()) > 10
            
            return extracted_text, True
        
        except Exception as e:
            logger.error(f"Error extracting text from DOCX: {e}")
            return "", False
    
    @staticmethod
    def extract_text_from_doc(file_path: str) -> Tuple[str, bool]:
        """
        Extract text from DOC file (using python-docx as fallback)
        For older DOC format, might need additional library like python-doc2docx
        Returns: (text, success)
        """
        try:
            # Try opening as DOCX first (newer Office formats)
            return CVParser.extract_text_from_docx(file_path)
        except Exception as e:
            # If python-docx fails, it's likely an older .doc format (pre-2007)
            logger.warning(f"Cannot parse .doc file - appears to be older MS Word format (pre-2007): {e}")
            logger.info("Supported formats: .pdf, .docx (Office 2007+)")
            return "", False
    
    @staticmethod
    def extract_text(file_path: str, file_extension: str) -> Tuple[str, bool]:
        """
        Extract text from file based on extension
        Returns: (text, success)
        """
        file_extension = file_extension.lower()
        
        if file_extension == "pdf":
            return CVParser.extract_text_from_pdf(file_path)
        elif file_extension == "docx":
            return CVParser.extract_text_from_docx(file_path)
        elif file_extension == "doc":
            return CVParser.extract_text_from_doc(file_path)
        else:
            logger.error(f"Unsupported file format: {file_extension}")
            return "", False
    
    @staticmethod
    def clean_text(text: str) -> str:
        """
        Clean extracted text
        Remove extra whitespace, normalize line breaks, etc.
        """
        # Remove extra whitespace
        text = " ".join(text.split())
        
        # Normalize line breaks
        text = text.replace("\\n", "\n")
        text = text.replace("\\t", "\t")
        
        # Remove special characters but keep structure
        text = text.replace("•", "-")
        text = text.replace("●", "-")
        text = text.replace("◦", "-")
        
        return text


# Helper function for quick parsing
def parse_cv_file(file_path: str, file_extension: str) -> Tuple[str, bool]:
    """
    Quick helper to parse CV file
    Returns: (text, success)
    """
    text, success = CVParser.extract_text(file_path, file_extension)
    if success:
        text = CVParser.clean_text(text)
    return text, success
