"""Generate the data structures Word document for MUQAYYIM."""
from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy

doc = Document()

# ── Page margins ──────────────────────────────────────────────────────────────
section = doc.sections[0]
section.top_margin    = Cm(2.5)
section.bottom_margin = Cm(2.5)
section.left_margin   = Cm(3)
section.right_margin  = Cm(2.5)

# ── Colour palette ────────────────────────────────────────────────────────────
# Tuples (R, G, B) used for XML hex strings; RGBColor used for font colours
INDIGO_T    = (0x4F, 0x46, 0xE5)
DARK_T      = (0x1E, 0x29, 0x3B)
MID_T       = (0x47, 0x55, 0x69)
LIGHT_BG_T  = (0xF1, 0xF5, 0xF9)
BORDER_T    = (0xE2, 0xE8, 0xF0)
WHITE_T     = (0xFF, 0xFF, 0xFF)
HEADER_BG_T = (0x4F, 0x46, 0xE5)

INDIGO    = RGBColor(*INDIGO_T)
DARK      = RGBColor(*DARK_T)
MID       = RGBColor(*MID_T)
WHITE     = RGBColor(*WHITE_T)

def _hex(t): return '{:02X}{:02X}{:02X}'.format(*t)

# ── Helper: set paragraph shading ────────────────────────────────────────────
def shade_cell(cell, rgb_tuple):
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd  = OxmlElement('w:shd')
    shd.set(qn('w:val'),   'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'),  _hex(rgb_tuple))
    tcPr.append(shd)

# ── Helper: set cell borders ──────────────────────────────────────────────────
def set_cell_border(cell, color_tuple=None):
    if color_tuple is None:
        color_tuple = BORDER_T
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')
    for side in ('top', 'left', 'bottom', 'right'):
        border = OxmlElement(f'w:{side}')
        border.set(qn('w:val'),   'single')
        border.set(qn('w:sz'),    '4')
        border.set(qn('w:space'), '0')
        border.set(qn('w:color'), _hex(color_tuple))
        tcBorders.append(border)
    tcPr.append(tcBorders)

# ── Helper: styled heading ────────────────────────────────────────────────────
def add_heading(doc, text, level=1):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run(text)
    run.bold = True
    if level == 1:
        run.font.size  = Pt(20)
        run.font.color.rgb = INDIGO
        pPr = p._p.get_or_add_pPr()
        pBdr = OxmlElement('w:pBdr')
        bottom = OxmlElement('w:bottom')
        bottom.set(qn('w:val'),   'single')
        bottom.set(qn('w:sz'),    '6')
        bottom.set(qn('w:space'), '4')
        bottom.set(qn('w:color'), '4F46E5')
        pBdr.append(bottom)
        pPr.append(pBdr)
        p.paragraph_format.space_before = Pt(18)
        p.paragraph_format.space_after  = Pt(8)
    elif level == 2:
        run.font.size  = Pt(14)
        run.font.color.rgb = DARK
        p.paragraph_format.space_before = Pt(14)
        p.paragraph_format.space_after  = Pt(4)
    else:
        run.font.size  = Pt(12)
        run.font.color.rgb = INDIGO
        p.paragraph_format.space_before = Pt(10)
        p.paragraph_format.space_after  = Pt(4)
    return p

# ── Helper: body paragraph ────────────────────────────────────────────────────
def add_body(doc, text, bold_prefix=None):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(5)
    if bold_prefix:
        r = p.add_run(bold_prefix + '  ')
        r.bold = True
        r.font.color.rgb = DARK
    r2 = p.add_run(text)
    r2.font.size      = Pt(11)
    r2.font.color.rgb = MID
    return p

# ── Helper: bullet ────────────────────────────────────────────────────────────
def add_bullet(doc, text):
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.space_after  = Pt(3)
    p.paragraph_format.left_indent  = Inches(0.3)
    run = p.add_run(text)
    run.font.size      = Pt(11)
    run.font.color.rgb = MID
    return p

# ── Helper: styled table ──────────────────────────────────────────────────────
def add_table(doc, headers, rows, col_widths=None):
    t = doc.add_table(rows=1 + len(rows), cols=len(headers))
    t.alignment = WD_TABLE_ALIGNMENT.LEFT
    t.style     = 'Table Grid'

    # Header row
    for i, h in enumerate(headers):
        cell = t.cell(0, i)
        shade_cell(cell, HEADER_BG_T)
        set_cell_border(cell, color_tuple=HEADER_BG_T)
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        p   = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        run = p.add_run(h)
        run.bold           = True
        run.font.size      = Pt(10)
        run.font.color.rgb = WHITE

    # Data rows
    for ri, row in enumerate(rows):
        bg = LIGHT_BG_T if ri % 2 == 0 else WHITE_T
        for ci, val in enumerate(row):
            cell = t.cell(ri + 1, ci)
            shade_cell(cell, bg)
            set_cell_border(cell, color_tuple=BORDER_T)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            p   = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            run = p.add_run(str(val))
            run.font.size      = Pt(10)
            run.font.color.rgb = MID

    # Column widths
    if col_widths:
        for ri in range(len(rows) + 1):
            for ci, w in enumerate(col_widths):
                t.cell(ri, ci).width = Inches(w)

    doc.add_paragraph()  # spacing after table
    return t

# ── Helper: code block ────────────────────────────────────────────────────────
def add_code(doc, lines):
    for line in lines:
        p   = doc.add_paragraph()
        p.paragraph_format.space_after  = Pt(1)
        p.paragraph_format.left_indent  = Inches(0.4)
        run = p.add_run(line)
        run.font.name      = 'Courier New'
        run.font.size      = Pt(9)
        run.font.color.rgb = RGBColor(0x1E, 0x40, 0xAF)  # blue code colour

# =============================================================================
#  COVER
# =============================================================================
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('MUQAYYIM')
r.bold = True; r.font.size = Pt(32); r.font.color.rgb = INDIGO
p.paragraph_format.space_before = Pt(48)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('AI-Powered CV Intelligence Platform')
r.font.size = Pt(16); r.font.color.rgb = MID

doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('Information Domain & Data Structures')
r.bold = True; r.font.size = Pt(18); r.font.color.rgb = DARK

doc.add_paragraph()
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('How real-world concepts are modelled, stored, processed, and organised')
r.font.size = Pt(11); r.font.color.rgb = MID

doc.add_page_break()

# =============================================================================
#  1. INTRODUCTION
# =============================================================================
add_heading(doc, '1. Introduction', 1)
add_body(doc,
    'MUQAYYIM is a microservices-based web application that automates the extraction and '
    'management of professional information from CVs. The system transforms unstructured '
    'real-world data — a PDF or DOCX résumé — into structured, queryable records that '
    'power a full professional profile, a GitHub-linked skills portfolio, and a dashboard '
    'that tracks the CV through its entire processing lifecycle.')

add_body(doc,
    'This document describes how each concept in the problem domain (a person, their career '
    'history, a CV file, a GitHub repository, a parsed skill) is represented as a concrete '
    'data structure inside the system, and how those structures flow across services.')

# =============================================================================
#  2. INFORMATION DOMAIN
# =============================================================================
add_heading(doc, '2. Information Domain Overview', 1)
add_body(doc,
    'The system deals with five core real-world concepts. Each concept maps to one or more '
    'data entities inside the application.')

add_table(doc,
    headers=['Domain Concept', 'System Entity', 'Storage Location'],
    rows=[
        ['A registered user',           'User document',          'MongoDB — users collection'],
        ['Professional profile',         'Profile document',       'MongoDB — profiles collection'],
        ['A CV file (PDF/DOCX)',         'Temporary file on disk', 'Server filesystem (uploads/)'],
        ['NLP-extracted CV data',        'CVParsedData document',  'MongoDB — cv_parsed_data collection'],
        ['GitHub repository metadata',   'Embedded in Profile',    'MongoDB — profiles collection (githubData)'],
    ],
    col_widths=[2.1, 2.1, 2.6]
)

add_body(doc,
    'The three MongoDB collections are the primary persistence layer. '
    'The temporary file on disk exists only between upload and verification; it is deleted '
    'once the user confirms their CV data.')

# =============================================================================
#  3. DATABASE
# =============================================================================
add_heading(doc, '3. Database & Storage Inventory', 1)

add_heading(doc, '3.1  Primary Database', 2)
add_table(doc,
    headers=['Property', 'Value'],
    rows=[
        ['Engine',      'MongoDB 6+'],
        ['Database name', 'muqayyim'],
        ['Interface',   'Mongoose (Node.js) / Motor async driver (Python/FastAPI)'],
        ['Connection',  'mongodb://localhost:27017/muqayyim'],
        ['Collections', 'users, profiles, cv_parsed_data'],
    ],
    col_widths=[2.2, 4.6]
)

add_heading(doc, '3.2  Ephemeral File Storage', 2)
add_table(doc,
    headers=['Property', 'Value'],
    rows=[
        ['Location',       'Module 2 backend — uploads/ directory'],
        ['Contents',       'Uploaded PDF/DOCX files during active parsing session'],
        ['Naming scheme',  '<uuid4>.<ext>   e.g.  f3a9b2c1-….pdf'],
        ['Lifetime',       'Created on POST /api/cv/upload; deleted on PUT /api/cv/verify'],
        ['Not persisted',  'Files are not backed up or replicated; only the extracted data is kept'],
    ],
    col_widths=[2.2, 4.6]
)

add_heading(doc, '3.3  In-Memory / Runtime Structures', 2)
add_table(doc,
    headers=['Structure', 'Where', 'Purpose'],
    rows=[
        ['NLPExtractor singleton',     'FastAPI process (Python)',    'Loaded once at startup; holds spaCy model + skill vocabulary sets'],
        ['JWT payload (decoded dict)', 'Every authenticated request', '{"userId", "id", "role"} — never persisted, lives in request scope only'],
        ['React AuthContext state',    'Browser (React)',             '{"user", "token"} — mirrors localStorage; refreshed on navigation'],
        ['localStorage (browser)',     'Client browser',             'Stores JWT string under key "token"; cleared on logout'],
    ],
    col_widths=[2.2, 2.1, 2.5]
)

# =============================================================================
#  4. ENTITY SCHEMAS
# =============================================================================
add_heading(doc, '4. Entity Schemas & Data Transformations', 1)

# ── 4.1 User ─────────────────────────────────────────────────────────────────
add_heading(doc, '4.1  User  (MongoDB — users collection)', 2)
add_body(doc,
    'The User entity represents a registered account. It is created when the user signs up '
    'and is the identity anchor for every other entity in the system. '
    'The password is never stored in plain text — it is bcrypt-hashed with a salt factor of 10 '
    'before the document is saved.')

add_table(doc,
    headers=['Field', 'Type', 'Constraint / Default', 'Purpose'],
    rows=[
        ['_id',              'ObjectId',  'Auto-generated',           'Primary key; referenced by Profile.userId'],
        ['name',             'String',    'Required, max 100 chars',  'Display name'],
        ['email',            'String',    'Unique, lowercase',        'Login credential'],
        ['password',         'String',    'Min 6 chars, select:false','bcrypt hash; never returned in API responses'],
        ['role',             'String',    'Enum: JobSeeker | Admin',  'Access control; embedded in JWT'],
        ['cvStatus',         'String',    'Enum: 4 values below',     'Tracks CV through the processing pipeline'],
        ['resetToken',       'String',    'Nullable',                 'SHA-256 hash of the emailed reset token'],
        ['resetTokenExpiry', 'Date',      'Nullable',                 '1-hour window after forgot-password request'],
        ['createdAt',        'Date',      'Auto (timestamps: true)',   'Account creation timestamp'],
        ['updatedAt',        'Date',      'Auto (timestamps: true)',   'Last modification timestamp'],
    ],
    col_widths=[1.6, 1.0, 1.8, 2.3]
)

add_body(doc, 'cvStatus state machine:', bold_prefix='')
add_table(doc,
    headers=['Value', 'Meaning', 'Set by'],
    rows=[
        ['Not Uploaded', 'No CV file has ever been uploaded',      'Default on registration'],
        ['Uploaded',     'File received; not yet parsed',           'POST /api/cv/upload succeeds'],
        ['Processing',   'NLP extraction is in progress / done',   'POST /api/cv/parse/:file_id succeeds'],
        ['Verified',     'User confirmed the extracted data',       'PUT /api/cv/verify succeeds'],
    ],
    col_widths=[1.4, 3.0, 2.4]
)

# ── 4.2 Profile ───────────────────────────────────────────────────────────────
add_heading(doc, '4.2  Profile  (MongoDB — profiles collection)', 2)
add_body(doc,
    'The Profile document holds the user\'s full professional portfolio. '
    'It is keyed on userId (a unique index) so there is at most one Profile per User. '
    'It is created or updated as a single upsert — missing fields are left as empty defaults '
    'rather than causing validation errors.')

add_body(doc, 'Top-level structure:')
add_table(doc,
    headers=['Field', 'Type', 'Notes'],
    rows=[
        ['_id',          'ObjectId', 'Auto-generated primary key'],
        ['userId',       'ObjectId', 'Ref → User._id  (unique index)'],
        ['personalInfo', 'Object',   'Embedded sub-document — see below'],
        ['summary',      'String',   'Free-text professional summary'],
        ['experience',   'Array',    'Array of experience sub-documents'],
        ['education',    'Array',    'Array of education sub-documents'],
        ['skills',       'Object',   'Two arrays: manual[] and extracted[]'],
        ['certifications','Array',   'Array of certification sub-documents'],
        ['projects',     'Object',   'Two arrays: manual[] and github[]'],
        ['githubData',   'Object',   'Raw GitHub fetch result + processed skills'],
        ['createdAt',    'Date',     'Auto (timestamps: true)'],
        ['updatedAt',    'Date',     'Auto (timestamps: true)'],
    ],
    col_widths=[1.6, 1.2, 3.9]
)

add_body(doc, 'personalInfo sub-document:')
add_table(doc,
    headers=['Field', 'Type', 'Default'],
    rows=[
        ['phone',    'String', '""'],
        ['location', 'String', '""'],
        ['website',  'String', '""'],
        ['linkedin', 'String', '""'],
        ['github',   'String', '""'],
    ],
    col_widths=[1.6, 1.2, 3.9]
)

add_body(doc, 'experience[] item:')
add_table(doc,
    headers=['Field', 'Type', 'Required', 'Notes'],
    rows=[
        ['company',     'String',  'Yes', ''],
        ['title',       'String',  'Yes', 'Job title'],
        ['location',    'String',  'No',  'Default ""'],
        ['startDate',   'String',  'No',  'e.g. "2022-01"'],
        ['endDate',     'String',  'No',  'Empty if current=true'],
        ['current',     'Boolean', 'No',  'Default false'],
        ['description', 'String',  'No',  'Role responsibilities'],
    ],
    col_widths=[1.4, 1.0, 0.9, 3.4]
)

add_body(doc, 'education[] item:')
add_table(doc,
    headers=['Field', 'Type', 'Required', 'Notes'],
    rows=[
        ['institution', 'String', 'Yes', ''],
        ['degree',      'String', 'No',  'e.g. "BSc Computer Science"'],
        ['field',       'String', 'No',  'Field of study'],
        ['startDate',   'String', 'No',  ''],
        ['endDate',     'String', 'No',  ''],
        ['gpa',         'String', 'No',  'Grade / GPA as free text'],
    ],
    col_widths=[1.4, 1.0, 0.9, 3.4]
)

add_body(doc, 'certifications[] item:')
add_table(doc,
    headers=['Field', 'Type', 'Required'],
    rows=[
        ['name',   'String', 'Yes'],
        ['issuer', 'String', 'No'],
        ['date',   'String', 'No'],
        ['url',    'String', 'No'],
    ],
    col_widths=[1.6, 1.2, 0.8]
)

add_body(doc, 'projects.manual[] item:')
add_table(doc,
    headers=['Field', 'Type', 'Required', 'Notes'],
    rows=[
        ['name',          'String',   'Yes', ''],
        ['description',   'String',   'No',  ''],
        ['url',           'String',   'No',  'Live URL or repo link'],
        ['technologies',  '[String]', 'No',  'Tag array e.g. ["React","MongoDB"]'],
    ],
    col_widths=[1.4, 1.0, 0.9, 3.4]
)

add_body(doc, 'projects.github[] item  (selected GitHub repos applied to profile):')
add_table(doc,
    headers=['Field', 'Type', 'Notes'],
    rows=[
        ['repoId',      'Number',   'GitHub numeric repository ID'],
        ['name',        'String',   'Repository name'],
        ['description', 'String',   ''],
        ['url',         'String',   'https://github.com/…'],
        ['languages',   '[String]', 'Programming languages used'],
        ['topics',      '[String]', 'GitHub topic tags'],
        ['stars',       'Number',   'Stargazer count'],
    ],
    col_widths=[1.4, 1.2, 4.1]
)

add_body(doc, 'githubData sub-document  (raw fetch cache):')
add_table(doc,
    headers=['Field', 'Type', 'Notes'],
    rows=[
        ['username',        'String',  'Last fetched GitHub username'],
        ['repos',           'Mixed',   'Raw array from GitHub API — all fetched repos'],
        ['extractedSkills', '[String]','Union of all repo languages + topics; deduplicated, sorted'],
        ['lastSynced',      'Date',    'Timestamp of last successful fetch'],
    ],
    col_widths=[1.6, 1.0, 4.1]
)

# ── 4.3 CVParsedData ──────────────────────────────────────────────────────────
add_heading(doc, '4.3  CVParsedData  (MongoDB — cv_parsed_data collection)', 2)
add_body(doc,
    'Managed exclusively by the AI Service (FastAPI / Python). '
    'One document is created per parsing session. A user can re-upload, '
    'which creates a new document; the summary endpoint always returns the most '
    'recent record sorted by upload_date descending.')

add_table(doc,
    headers=['Field', 'Type', 'Notes'],
    rows=[
        ['_id',            'ObjectId', 'Auto-generated'],
        ['user_id',        'String',   'Extracted from JWT — never from request body'],
        ['file_name',      'String',   '"<uuid>.<ext>" — matches the filename on disk'],
        ['file_path',      'String',   'Absolute path on the AI Service host'],
        ['upload_date',    'DateTime', 'UTC; set at creation time'],
        ['parsed_data',    'Object',   'NLP output — see ParsedDataModel below'],
        ['parsing_status', 'String',   '"completed" → "verified" after user confirms'],
    ],
    col_widths=[1.6, 1.2, 3.9]
)

add_body(doc, 'ParsedDataModel (embedded):')
add_table(doc,
    headers=['Sub-field', 'Type', 'Item fields'],
    rows=[
        ['skills[]',    'Array', 'name: String, confidence: Float (0–1)'],
        ['education[]', 'Array', 'degree, institution, year, confidence'],
        ['experience[]','Array', 'title, company, duration, confidence'],
    ],
    col_widths=[1.4, 0.9, 4.4]
)

add_body(doc,
    'confidence is a float between 0 and 1 assigned by the NLP extractor. '
    'Values ≥ 0.95 indicate a direct keyword match from the curated skills database; '
    '0.85 indicates a full-text match; 0.65–0.75 indicates a contextual or heuristic match.')

# =============================================================================
#  5. DATA TRANSFORMATION PIPELINE
# =============================================================================
add_heading(doc, '5. Data Transformation Pipeline', 1)
add_body(doc,
    'The transformation from an unstructured CV file to verified structured data '
    'passes through five distinct stages. Each stage changes the form of the data.')

add_table(doc,
    headers=['Stage', 'Input', 'Process', 'Output', 'Storage'],
    rows=[
        ['1 — Upload',
         'Binary file stream (PDF/DOCX)',
         'Validate MIME type and size (≤5 MB); generate UUID filename',
         'File on disk + file_id string returned to client',
         'uploads/ directory'],
        ['2 — Text Extraction',
         'File path + extension',
         'pdfplumber (PDF) or python-docx (DOCX) reads raw bytes → plain UTF-8 text',
         'cv_text: str',
         'In-memory only'],
        ['3 — NLP Extraction',
         'cv_text: str',
         'spaCy en_core_web_sm NER + regex section splitter + curated keyword vocabulary',
         'ParsedDataModel {skills[], education[], experience[]}',
         'Inserted into cv_parsed_data collection'],
        ['4 — User Verification',
         'ParsedDataModel (browser display)',
         'User reviews, edits, and confirms extracted items',
         'Verified ParsedDataModel written back to cv_parsed_data; cvStatus → Verified',
         'MongoDB cv_parsed_data; User.cvStatus updated'],
        ['5 — Profile Save',
         'ProfileBuilderPage form state',
         'User manually adds/edits fields; GitHub fetch merges repo data',
         'Profile document upserted in MongoDB',
         'MongoDB profiles collection'],
    ],
    col_widths=[1.1, 1.4, 1.9, 1.6, 1.2]
)

# =============================================================================
#  6. NLP EXTRACTION DETAIL
# =============================================================================
add_heading(doc, '6. NLP Extraction — Internal Data Structures', 1)
add_body(doc,
    'The NLPExtractor class (Module 2, nlp_extractor.py) holds three in-memory data '
    'structures that are loaded once at process startup and shared across all requests '
    'via a module-level singleton.')

add_table(doc,
    headers=['Structure', 'Python Type', 'Size / Contents', 'Role'],
    rows=[
        ['technical_skills',
         'set[str]',
         '~120 entries — programming languages, frameworks, databases, cloud tools',
         'Fast O(1) membership test for skill keyword matching'],
        ['soft_skills',
         'set[str]',
         '~12 entries — leadership, communication, teamwork, etc.',
         'Same; lower confidence score assigned (0.85 vs 0.95)'],
        ['degree_keywords',
         'set[str]',
         '~10 entries — bachelor, master, phd, diploma, etc.',
         'Triggers education extraction; without a match the line is skipped'],
        ['nlp (spaCy model)',
         'spacy.Language',
         '~12 MB on disk — en_core_web_sm',
         'Named Entity Recognition for ORG entities (institution/company names)'],
    ],
    col_widths=[1.6, 1.1, 2.2, 2.0]
)

add_body(doc, 'Section-splitting algorithm:')
add_bullet(doc, 'The raw CV text is scanned for heading keywords (Education, Skills, Experience, etc.) using a compiled regex.')
add_bullet(doc, 'Each heading marks the start of a section; the text between consecutive headings is treated as that section\'s content.')
add_bullet(doc, 'Skill extraction targets the Skills section first; falls back to full-text scan if no Skills heading is found.')
add_bullet(doc, 'Experience extraction targets Work Experience / Professional Experience sections and processes line-by-line.')
add_bullet(doc, 'Results are deduplicated using Python dicts keyed on normalised skill/entity names.')

# =============================================================================
#  7. GITHUB INTEGRATION
# =============================================================================
add_heading(doc, '7. GitHub Integration — Data Flow', 1)
add_body(doc,
    'The githubService.js module (Core Service) transforms a GitHub username into structured '
    'skill and project data through the following steps.')

add_table(doc,
    headers=['Step', 'Action', 'Data Shape'],
    rows=[
        ['1 — API Call',
         'GET https://api.github.com/users/{username}/repos?per_page=100&sort=updated',
         'Raw JSON array of up to 100 repository objects'],
        ['2 — Mapping',
         'Map each raw repo to a clean internal object',
         '{ repoId, name, description, url, language, topics[], stars, fork, updatedAt }'],
        ['3 — Skill Extraction',
         'Collect repo.language (single string) and repo.topics (string array) from all repos',
         'Flat array with duplicates'],
        ['4 — Deduplication',
         'Pass through Set; filter falsy values; sort alphabetically',
         'extractedSkills: String[]  — e.g. ["JavaScript","React","nodejs","python"]'],
        ['5 — Storage',
         'Saved to Profile.githubData via findOneAndUpdate upsert',
         'githubData.{ username, repos[], extractedSkills[], lastSynced }'],
        ['6 — User Selection',
         'User selects repos in the frontend; clicks Apply',
         'Selected repos → Profile.projects.github[];  skills → Profile.skills.extracted[]'],
    ],
    col_widths=[0.8, 2.5, 3.4]
)

add_body(doc,
    'Authentication: if the GITHUB_TOKEN environment variable is set, every request to the '
    'GitHub API includes an Authorization: Bearer header, raising the rate limit from 60 to '
    '5000 requests per hour. Without it the service still works for public repos within the '
    'unauthenticated quota.')

# =============================================================================
#  8. JWT & IDENTITY
# =============================================================================
add_heading(doc, '8. JWT & Identity Data Structure', 1)
add_body(doc,
    'Authentication state is carried entirely inside a signed JWT. '
    'No session table or server-side session store exists. '
    'The token is the sole shared artefact between Module 1 (issuer) and Module 2 (consumer).')

add_body(doc, 'JWT payload structure:')
add_code(doc, [
    '{',
    '  "userId" : "64f3a9b2c1d4e5f6a7b8c9d0",   // MongoDB ObjectId string',
    '  "id"     : "64f3a9b2c1d4e5f6a7b8c9d0",   // alias for userId',
    '  "role"   : "JobSeeker",',
    '  "iat"    : 1700000000,',
    '  "exp"    : 1700604800                      // 7-day expiry',
    '}',
])
doc.add_paragraph()

add_table(doc,
    headers=['Layer', 'Storage', 'Key name', 'Action on login', 'Action on logout'],
    rows=[
        ['Core Service (Node)', 'Not stored', '—', 'Signs JWT with JWT_SECRET env var', '—'],
        ['Browser (React)', 'localStorage', '"token"', 'Written by authService.login()', 'Cleared by authService.logout()'],
        ['AI Service (Python)', 'Not stored', '—', 'Reads Authorization header; verifies with same JWT_SECRET', '—'],
    ],
    col_widths=[1.5, 1.2, 1.0, 2.0, 1.5]
)

# =============================================================================
#  9. COLLECTION RELATIONSHIPS
# =============================================================================
add_heading(doc, '9. Collection Relationships', 1)
add_body(doc,
    'MongoDB does not enforce foreign keys. The relationships below are maintained '
    'by application code (Mongoose schemas and service logic).')

add_table(doc,
    headers=['Relationship', 'Type', 'Enforced by'],
    rows=[
        ['User  ──►  Profile',       '1 : 0..1  (unique index on userId)',  'Profile.userId has unique: true index'],
        ['User  ──►  CVParsedData',  '1 : many  (user_id string field)',     'Application query; no DB constraint'],
        ['Profile  ─┐',             '',                                       ''],
        ['  ├  experience[]',        'Embedded array',                        'Mongoose sub-schema inside Profile'],
        ['  ├  education[]',         'Embedded array',                        'Mongoose sub-schema inside Profile'],
        ['  ├  certifications[]',    'Embedded array',                        'Mongoose sub-schema inside Profile'],
        ['  ├  projects.manual[]',   'Embedded array',                        'Mongoose sub-schema inside Profile'],
        ['  ├  projects.github[]',   'Embedded array',                        'Populated by GitHub fetch + user selection'],
        ['  └  githubData',          'Embedded object',                       'Overwritten on each GitHub sync'],
    ],
    col_widths=[2.1, 2.1, 2.6]
)

add_body(doc,
    'All arrays (experience, education, certifications, projects) are embedded subdocuments '
    'rather than separate collections. This is deliberate: a profile and all its sub-items '
    'are always fetched together and always belong to one user, so embedding eliminates joins '
    'and keeps the read path to a single document lookup.')

# =============================================================================
#  10. API GATEWAY ROUTING TABLE
# =============================================================================
add_heading(doc, '10. API Gateway — Routing & Data Flow', 1)
add_body(doc,
    'The API Gateway (port 8080) is a thin Express proxy. It does not transform data; '
    'it only routes HTTP requests to the correct backend service and forwards the '
    'Authorization header unchanged.')

add_table(doc,
    headers=['Route prefix', 'Proxied to', 'Port', 'Data handled'],
    rows=[
        ['/api/auth/*',    'Core Service',  '3000', 'User registration, login, JWT, CV status, account info'],
        ['/api/profile/*', 'Core Service',  '3000', 'Professional profile CRUD, GitHub fetch'],
        ['/api/cv/*',      'AI Service',    '8000', 'CV file upload, NLP parse, verify, summary'],
    ],
    col_widths=[1.7, 1.5, 0.7, 2.8]
)

# =============================================================================
#  11. DATA LIFECYCLE SUMMARY
# =============================================================================
add_heading(doc, '11. End-to-End Data Lifecycle', 1)

add_heading(doc, '11.1  User Registration', 3)
add_bullet(doc, 'Client sends {name, email, password, passwordConfirm} to POST /api/auth/register via gateway.')
add_bullet(doc, 'Core Service validates; bcrypt hashes password; creates User document; signs JWT.')
add_bullet(doc, 'Response {token, user} is written to browser localStorage. Profile does not exist yet.')

add_heading(doc, '11.2  CV Upload → Parse → Verify', 3)
add_bullet(doc, 'Client POSTs multipart/form-data file to /api/cv/upload. AI Service saves file to disk; returns file_id.')
add_bullet(doc, 'Core Service sets User.cvStatus = "Uploaded".')
add_bullet(doc, 'Client POSTs to /api/cv/parse/{file_id}. AI Service reads file, runs NLP, inserts CVParsedData document.')
add_bullet(doc, 'Core Service sets User.cvStatus = "Processing".')
add_bullet(doc, 'Client reviews extracted data, edits if needed, then PUTs to /api/cv/verify with confirmed data.')
add_bullet(doc, 'AI Service updates the CVParsedData document (parsing_status → "verified"); deletes the temp file from disk.')
add_bullet(doc, 'Core Service sets User.cvStatus = "Verified".')

add_heading(doc, '11.3  Profile Build & GitHub Sync', 3)
add_bullet(doc, 'User fills in profile form sections; React state is collected under a single profile object.')
add_bullet(doc, 'User enters GitHub username → POSTs to /api/profile/github → Core Service calls GitHub API → returns repos + extractedSkills.')
add_bullet(doc, 'User selects repos in the UI; clicks Apply → selected repos merged into profile.projects.github; skills into profile.skills.extracted.')
add_bullet(doc, 'User clicks Save Profile → PUTs to /api/profile with the full profile object → Core Service upserts Profile document.')

# =============================================================================
#  FOOTER note
# =============================================================================
doc.add_paragraph()
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('MUQAYYIM — Final Year Project  ·  Data Structures & Information Domain')
r.font.size = Pt(9)
r.font.color.rgb = RGBColor(0x94, 0xA3, 0xB8)  # slate footer
r.italic = True

# ── Save ──────────────────────────────────────────────────────────────────────
out = r"d:\FYP\Module 1,2\MUQAYYIM_Data_Structures.docx"
doc.save(out)
print(f"Saved: {out}")
